#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
翻译合同为中英文对照格式
每个段落：中文在上，英文在下
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googletrans import Translator
import time

def translate_text(text, translator, retry_count=3):
    """使用 Google 翻译中文到英文"""
    if not text or not text.strip():
        return ""

    for attempt in range(retry_count):
        try:
            result = translator.translate(text, src='zh-CN', dest='en')
            return result.text
        except Exception as e:
            print(f"Translation error (attempt {attempt + 1}): {e}")
            if attempt < retry_count - 1:
                time.sleep(2)  # 等待后重试
            else:
                return f"[Translation Error: {str(e)}]"

def process_contract(input_path, output_path):
    """处理合同文件"""
    # 创建翻译器
    translator = Translator(service_urls=['translate.googleapis.com'])

    # 读取原文件
    print(f"Reading {input_path}...")
    doc = Document(input_path)

    # 收集需要翻译的段落数量和内容
    paragraphs_to_translate = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 1:  # 忽略空段落和单个字符
            paragraphs_to_translate.append({
                'index': i,
                'text': text,
                'style': para.style.name,
                'alignment': para.alignment
            })

    print(f"Found {len(paragraphs_to_translate)} paragraphs to translate")

    # 批量翻译
    print("Starting translation...")
    translated_texts = {}
    for idx, item in enumerate(paragraphs_to_translate, 1):
        print(f"[{idx}/{len(paragraphs_to_translate)}] Translating: {item['text'][:50]}...")
        translated = translate_text(item['text'], translator)
        translated_texts[item['index']] = translated

        # 避免触发速率限制
        if idx < len(paragraphs_to_translate):
            time.sleep(0.5)

    print("Translation completed!")

    # 创建新文档
    print(f"Creating translated document: {output_path}")
    new_doc = Document()

    # 标题
    title = new_doc.add_paragraph("青岛九瑞区域合作协议（中英对照版）")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    new_doc.add_paragraph()  # 空行

    # 添加中英文段落
    for item in paragraphs_to_translate:
        # 中文段落
        cn_para = new_doc.add_paragraph(item['text'])
        cn_run = cn_para.runs[0]
        cn_run.font.size = Pt(12)
        cn_run.font.name = '宋体'
        if item['alignment']:
            cn_para.alignment = item['alignment']

        # 英文段落
        en_text = translated_texts.get(item['index'], "")
        if en_text:
            en_para = new_doc.add_paragraph(en_text)
            en_run = en_para.runs[0]
            en_run.font.size = Pt(10)
            en_run.font.name = 'Times New Roman'
            en_run.font.color.rgb = RGBColor(80, 80, 80)  # 灰色区分
            if item['alignment']:
                en_para.alignment = item['alignment']

        new_doc.add_paragraph()  # 段落间距

    # 保存文档
    new_doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return True

if __name__ == "__main__":
    input_file = "/mnt/c/Users/Administrator/Desktop/青岛九瑞区域合作协议.docx"
    output_file = "/mnt/c/Users/Administrator/Desktop/青岛九瑞区域合作协议_中英对照.docx"

    import os
    if os.path.exists(input_file):
        process_contract(input_file, output_file)
    else:
        print(f"Error: Input file not found: {input_file}")
